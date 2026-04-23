"""
Local Ingestion — process .docx and .xlsx files from a local folder.
Chunks, embeds, and inserts into Supabase.

Resumable: skips files already logged in local_ingest_progress.json.

Run standalone:  python local_ingest.py --path ~/path/to/folder
Via unified CLI: python ingest.py --source local --path ~/path/to/folder
"""

import os
import json
import time
from pathlib import Path
from dotenv import load_dotenv
from google import genai
from google.genai import types
from supabase import create_client
from docx import Document as DocxDocument
from openpyxl import load_workbook

# ─── Config ────────────────────────────────────────────────────────────
load_dotenv()

GCP_PROJECT_ID = os.getenv("GCP_PROJECT_ID")
GCP_LOCATION = os.getenv("GCP_LOCATION")
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY")

PROGRESS_FILE = "local_ingest_progress.json"

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
MIN_CHUNK_SIZE = 100
EMBED_BATCH_SIZE = 100


# ─── File readers ──────────────────────────────────────────────────────
def read_docx(path):
    """Extract all paragraph text from a .docx file."""
    try:
        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n\n".join(paragraphs)
    except Exception as e:
        raise RuntimeError(f"Cannot read docx: {e}")


def read_xlsx(path):
    """Extract text from all cells in all sheets of an .xlsx file."""
    try:
        wb = load_workbook(path, data_only=True)
        lines = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            lines.append(f"=== Sheet: {sheet} ===")
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(c) for c in row if c is not None)
                if row_text.strip():
                    lines.append(row_text)
        return "\n".join(lines)
    except Exception as e:
        raise RuntimeError(f"Cannot read xlsx: {e}")


# ─── Chunking ──────────────────────────────────────────────────────────
def chunk_text(text, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    text = text.strip()
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + size
        chunk = text[start:end]
        if len(chunk) >= MIN_CHUNK_SIZE:
            chunks.append(chunk)
        start += size - overlap
    return chunks


# ─── Embedding ─────────────────────────────────────────────────────────
def embed_batch(client, texts):
    result = client.models.embed_content(
        model="gemini-embedding-001",
        contents=texts,
        config=types.EmbedContentConfig(task_type="RETRIEVAL_DOCUMENT"),
    )
    return [e.values for e in result.embeddings]


# ─── Progress tracking ─────────────────────────────────────────────────
def load_progress():
    if os.path.exists(PROGRESS_FILE):
        with open(PROGRESS_FILE) as f:
            return json.load(f)
    return {"completed_paths": []}


def save_progress(progress):
    with open(PROGRESS_FILE, "w") as f:
        json.dump(progress, f, indent=2)


# ─── Core ingestion ────────────────────────────────────────────────────
def run_local_ingest(genai_client, supabase, source_folder):
    """Run local file ingestion using pre-initialized clients. Called by ingest.py."""
    source_folder = Path(source_folder).expanduser().resolve()

    print("=" * 60)
    print("ASK-YOU — LOCAL INGESTION")
    print("=" * 60)

    if not source_folder.exists():
        print(f"\nERROR: Folder not found: {source_folder}")
        return

    print(f"\nScanning {source_folder}...")
    docx_files = list(source_folder.rglob("*.docx"))
    xlsx_files = list(source_folder.rglob("*.xlsx"))
    all_files = sorted(docx_files + xlsx_files, key=lambda p: str(p))
    print(f"  Found {len(docx_files)} .docx files")
    print(f"  Found {len(xlsx_files)} .xlsx files")
    print(f"  Total: {len(all_files)} files to process")

    progress = load_progress()
    done = set(progress["completed_paths"])
    remaining = [f for f in all_files if str(f) not in done]
    print(f"\nAlready completed: {len(done)}")
    print(f"Remaining: {len(remaining)}")

    if not remaining:
        print("\nAll files already ingested.")
        return

    total_chunks = 0
    failed = []
    start = time.time()

    for i, path in enumerate(remaining, 1):
        name = path.name
        print(f"\n[{i}/{len(remaining)}] {name}")

        try:
            if path.suffix.lower() == ".docx":
                text = read_docx(path)
            elif path.suffix.lower() == ".xlsx":
                text = read_xlsx(path)
            else:
                continue

            chunks = chunk_text(text)
            if not chunks:
                print("  (skipped — empty or too short)")
                progress["completed_paths"].append(str(path))
                save_progress(progress)
                continue

            print(f"  {len(chunks)} chunks, {len(text)} chars")

            vectors = []
            for b in range(0, len(chunks), EMBED_BATCH_SIZE):
                batch = chunks[b:b + EMBED_BATCH_SIZE]
                vectors.extend(embed_batch(genai_client, batch))

            rows = [
                {
                    "content": chunk,
                    "embedding": vec,
                    "source_doc": name,
                    "chunk_index": idx,
                }
                for idx, (chunk, vec) in enumerate(zip(chunks, vectors))
            ]

            INSERT_BATCH = 50
            for b in range(0, len(rows), INSERT_BATCH):
                supabase.table("knowledge").insert(rows[b:b + INSERT_BATCH]).execute()

            total_chunks += len(rows)
            print(f"  ✓ Inserted {len(rows)} chunks")

            progress["completed_paths"].append(str(path))
            save_progress(progress)

        except Exception as e:
            print(f"  ✗ FAILED: {e}")
            failed.append({"path": str(path), "name": name, "error": str(e)})

    elapsed = time.time() - start
    print("\n" + "=" * 60)
    print("INGESTION COMPLETE")
    print("=" * 60)
    print(f"Files processed:   {len(remaining) - len(failed)} / {len(remaining)}")
    print(f"Chunks inserted:   {total_chunks}")
    print(f"Failed files:      {len(failed)}")
    print(f"Elapsed:           {elapsed/60:.1f} minutes")

    if failed:
        with open("local_ingest_failures.json", "w") as f:
            json.dump(failed, f, indent=2)
        print("\nFailures logged to local_ingest_failures.json")


# ─── Standalone entry point ────────────────────────────────────────────
def main():
    import argparse
    parser = argparse.ArgumentParser(description="Ingest local .docx/.xlsx files into Supabase")
    parser.add_argument("--path", required=True, help="Folder containing files to ingest")
    args = parser.parse_args()

    print("\nInitializing Vertex AI and Supabase...")
    genai_client = genai.Client(vertexai=True, project=GCP_PROJECT_ID, location=GCP_LOCATION)
    supabase = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)
    print("  ✓ Ready")

    run_local_ingest(genai_client, supabase, args.path)


if __name__ == "__main__":
    main()
